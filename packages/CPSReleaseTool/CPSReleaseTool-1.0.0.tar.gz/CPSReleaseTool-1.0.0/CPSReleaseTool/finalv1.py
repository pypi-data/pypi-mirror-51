import tkinter as tk
from tkinter import messagebox
from tkinter import *
from tkinter.ttk import *
from tkinter import filedialog
from tkinter.filedialog import askopenfilename
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from tkinter import HORIZONTAL
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
import openpyxl
from docx.dml.color import ColorFormat
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import requests
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter,A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus.tables import Table
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.platypus.tables import Table
from reportlab import platypus
from  reportlab.lib.styles import ParagraphStyle as PS
import sys
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont      
import reportlab
from reportlab.lib.units import mm
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
import threading

try:

    class progressBar:
         
        def __init__(self):
            self.top = Tk()
            self.top.geometry("400x50")  
            self.frame = Frame(self.top)
            self.frame.pack()

            self.progress = Progressbar(self.top, orient = HORIZONTAL,  length = 350, mode = 'determinate')
            self.progress.pack(pady = 10)

            self.top.mainloop()

        def progBar(self):
            self.top.destroy()
            
   

    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row
        
        
            
    def addPageNumber(canvas, doc):
        """
        Add the page number
        """
        page_num = canvas.getPageNumber()
        text = "%s" % page_num
        canvas.drawRightString(200*mm, 20*mm, text)



    def add_hyperlink(paragraph, text, url): #used to add hyperlink to a particular cell text in table(docx).
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink


    #def getURLS():
     #   wb = openpyxl.load_workbook("G:\python\SPR Task Report.xlsx")
      #  sheet = wb["Sheet1"]
       # cells=sheet["E2":"F"+str(sheet.max_row)] #selecting the first cell below the heading row upto the last filled cell of the excel sheet
       # heading=sheet['A1':'C1']

        #for c1, c2 in cells:
         #   SPRLink="https://support.ptc.com/appserver/common/login/ssl/login.jsp?dest=%2Fappserver%2Fcs%2Fview%2Fspr.jsp%3Fn%3D"+str(c1.value)+"&msg=1" #dynamic link for article no. by just chaning the some letters of the link.
          #  ArticleLink="https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c2.value)

           # sprlinklist=[SPRLink]
            #articlelinklist=[ArticleLink]

            #for sl in sprlinklist:
             #   print(sl)

            #for al in articlelinklist:
             #   print(al)

    #this function is under testing..
    def fetchCells(filename,sheetname,savedfile):
        wb = openpyxl.load_workbook(filename)
        sheet = wb[sheetname]

        sheet.auto_filter.ref=sheet.dimensions
        
        cells=sheet["G2":"R"+str(sheet.max_row)] #selecting the first cell below the heading row upto the last filled cell of the excel sheet
        
        for c1,c2,c3,c4,c5,c6,c7,c8,c9,c10,c11,c12 in cells:
            if((c1.value=='Closed') and (c5.value=='Fix Submitted') and (str(c9.value)=="Yes") and (str(c10.value)=="True")): #add conditions for include rel notes and is customer spr C1 indicates col1 and so on...
                
                r  = requests.get("https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c11.value),headers={"User-Agent":"Mozilla/5.0"})
                data = r.text #getting the whole data from the link above.
                soup = BeautifulSoup(data, 'lxml') #parsing the data from the link.
                articlelink="https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c12.value) #dynamic link for spr no. by just changing the last 6 letters of the link.
                sprlink="https://support.ptc.com/appserver/common/login/ssl/login.jsp?dest=%2Fappserver%2Fcs%2Fview%2Fspr.jsp%3Fn%3D"+str(c11.value)+"&msg=1" #dynamic link for article no. by just changing the some letters of the link.

                sprlinklist=[sprlink]
                articlelinklist=[articlelink]
            
                print("\n Article:"+str(articlelinklist))
                print("\n SPR:"+str(sprlinklist))

            
        #under testing..
        cells1=cells[0:sheet.max_row//4]
        cells2=cells[sheet.max_row//4:sheet.max_row//2] 
        cells3=cells[sheet.max_row//2:sheet.max_row] 
        
        #print("\n \n cells1"+str(cells1))
        #print("\n \n cells2"+str(cells2))
        #print("\n \n cells3"+str(cells3))

        executor = ThreadPoolExecutor(max_workers=4)
        
        #task1 = executor.submit(toWord(savedfile,cells1,sheetname,sheet,wb))
        #task2 = executor.submit(toWord(savedfile,cells2,sheetname,sheet,wb))
        #task3 = executor.submit(toWord(savedfile,cells3,sheetname,sheet,wb))
    
        task1 = executor.submit(toWord(savedfile,cells,sheetname,sheet,wb))


    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row
        
        
    #def toWord(savedfile,cells,sheetname,sheet,wb):
    def toWord(filename,releaseno):
        start_time = datetime.now()

        wb = openpyxl.load_workbook(filename)
        sheet = wb.active

        sheet.auto_filter.ref=sheet.dimensions

        cells=sheet["G2":"R"+str(sheet.max_row)] #selecting the first cell below the heading row upto the last filled cell of the excel sheet

        document = Document()
        styles = document.styles

        phead=document.add_heading(releaseno, level=1) #heading in word file.
        run7 = phead.runs
        font7 = run7[0].font
        font7.name = 'Arial'
        font7.size = Pt(15.94)
        font7.bold = True
        font7.color.rgb = RGBColor(0, 0, 0)
        

        sprhead=document.add_heading('SPR Fixes', level=1)
        run7 = sprhead.runs
        font7 = run7[0].font
        font7.name = 'Arial'
        font7.size = Pt(11.96)
        font7.bold = True
        font7.color.rgb = RGBColor(0, 0, 0)

        
        sprpara=document.add_paragraph('This section lists the server-side SPRs and the Windchill Workgroup Manager SPRs that were fixed.')
        run7 = sprpara.runs
        font7 = run7[0].font
        font7.name = 'Times New Roman'
        font7.size = Pt(11.96)
        font7.color.rgb = RGBColor(0, 0, 0)

        
        serversprhead=document.add_heading('Server-Side SPRs', level=1)
        run7 = serversprhead.runs
        font7 = run7[0].font
        font7.name = 'Arial'
        font7.size = Pt(9.96)
        font7.bold = True
        font7.color.rgb = RGBColor(0, 0, 0)

        serverspr=document.add_paragraph('This section lists the server-side SPRs that were fixed.')    
        run7 = serverspr.runs
        font7 = run7[0].font
        font7.name = 'Times New Roman'
        font7.size = Pt(11.96)
        font7.color.rgb = RGBColor(0, 0, 0)
  
        table = document.add_table(rows=1, cols=3,style="Table Grid") #making a table with one row and 3 columns for heading row for first time.

        table.allow_autofit=True

        hdr_cells = table.rows[0].cells #heading cells.
        hdr_cells[0].width = Inches(1.5)
        hdr_cells[1].width = Inches(1.5)
        hdr_cells[2].width = Inches(22)
       
        hdr_cells[0].text ='SPR'
        hdr_cells[1].text ='Article'
        hdr_cells[2].text ='Description'

        set_repeat_table_header(table.rows[0])

        
        #for getting data from links using excel columns
        for c1,c2,c3,c4,c5,c6,c7,c8,c9,c10,c11,c12 in cells:
            
            if((c1.value=='Closed') and (c5.value=='Fix Submitted') and (str(c9.value)=="Yes") and (str(c10.value)=="True")): #add conditions for include rel notes and is customer spr C1 indicates col1 and so on...
                row_cells = table.add_row().cells #add a new row as the data comes in from the excel dynamically.
               
                r  = requests.get("https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c11.value),headers={"User-Agent":"Mozilla/5.0"})
                data = r.text #getting the whole data from the link above.
                soup = BeautifulSoup(data, 'lxml') #parsing the data from the link.
                sprlink="https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c11.value) #dynamic link for spr no. by just changing the last 6 letters of the link.
                articlelink="https://support.ptc.com/appserver/common/login/ssl/login.jsp?dest=%2Fappserver%2Fcs%2Fview%2Fspr.jsp%3Fn%3D"+str(c12.value)+"&msg=1" #dynamic link for article no. by just changing the some letters of the link.


                sprlinklist=[sprlink]
                articlelinklist=[articlelink]


                # for contents of the table
            
                for h1 in soup.findAll('h1',{"class":'article_solution_title_h1'}): #finding all the elements with h1 tag and article_solution_title_h1 class in the web page to get the content para.
                    
                    row_cells[2].text = (h1.get_text()).strip()  #setting the 3rd col of table to the data fetched from the page with h1 tag
                    paragraphdesc =row_cells[2].paragraphs[0] #used for styling the normal cell contents of spr col.

                    run6 =   paragraphdesc.runs
                    font6 =  run6[0].font
                    font6.name = 'Arial'
                    font6.size= Pt(11.96)
                    font6.color.rgb = RGBColor(0, 0, 0)


                    p = row_cells[0].add_paragraph()  
                    add_hyperlink(p, (str(c12.value)).strip(), articlelink) #setting the 1st col of table to the spr no. from excel sheet

                    pformat = p.paragraph_format
                    pformat.space_after = Pt(10)
                    
                    run4 =   p.runs
                    font4 =  run4[0].font
                    font4.name = 'Arial'
                    font4.size= Pt(11.96)
                    font4.color.rgb = RGBColor(0, 0, 0)

                    p2 = row_cells[1].add_paragraph()
                    add_hyperlink(p2,(str(c11.value)).strip(), sprlink) #setting the 1st col of table to the article no. from excel sheet

                    p2format = p2.paragraph_format
                    p2format.space_after = Pt(10)
                    
                    run5 =   p2.runs
                    font5 =  run5[0].font
                    font5.name = 'Arial'
                    font5.size= Pt(11.96)
                    font5.color.rgb = RGBColor(0, 0, 0)

                for header1 in soup.findAll('h1',{"class":'article_solution_title_h1'}):
                    for rowNum in range(2, sheet.max_row): #first row skip as it is heading
                       sheet.cell(row=rowNum,column=3).value = header1.get_text()

                    

                    #styling the heading or the first row of the table
                    paragraph1 = hdr_cells[0].paragraphs[0] #used for
                    paragraph2 = hdr_cells[1].paragraphs[0] #styling the
                    paragraph3 = hdr_cells[2].paragraphs[0] #heading cell contents.
                   
                    run1 =   paragraph1.runs
                    run2 =   paragraph2.runs
                    run3 =   paragraph3.runs

                    font1 =  run1[0].font
                    font2 =  run2[0].font
                    font3 =  run3[0].font

                    font1.name = 'Arial'
                    font2.name = 'Arial'
                    font3.name = 'Arial'

                    font1.size= Pt(11.96)
                    font2.size= Pt(11.96)
                    font3.size= Pt(11.96)

                    font1.bold = True
                    font2.bold = True
                    font3.bold = True

        document.add_page_break()

        document.save(releaseno+'.docx')
        wb.close()
        endtime=datetime.now()

        print('word file generated by name:'+releaseno+'.docx')
        messagebox.showinfo("Word file Generated!", 'word file generated by name:'+releaseno+'.docx')
        print(f'Time taken : {endtime-start_time}')





    def exceltopdf(filename,releaseno):

        start_time = datetime.now()

        wb = openpyxl.load_workbook(filename)
        sheet = wb.active

        sheet.auto_filter.ref=sheet.dimensions

        cells=sheet["G2":"R"+str(sheet.max_row)] #selecting the first cell below the heading row upto the last filled cell of the excel sheet

        cnt=0

        pdfmetrics.registerFont(TTFont('Arial', 'arialbd.ttf'))
        
        doc = SimpleDocTemplate(releaseno+".pdf", pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()

        elements.append(Paragraph('<para spaceafter="10" align="LEFT"><font color="black" size="15.94">'+releaseno+'</font></para>', styles['Heading1']))

        elements.append(Paragraph('<para align="LEFT"><font color="black" size="11.96">SPR Fixes</font></para>', styles['Heading2']))

        elements.append(Paragraph('<para spaceafter="20" align="LEFT"><font color="black" size="11.96">This section lists the server-side SPRs and the Windchill Workgroup Manager SPRs that fixed.</font></para>', styles["Normal"]))

        elements.append(Paragraph('<para align="LEFT"><font color="black" size="9.96">Server-Side SPRs</font></para>', styles['Heading2']))
        
        elements.append(Paragraph('<para spaceafter="5" align="LEFT"><font color="black" size="11.96">This section lists the server-side SPRs that were fixed.</font></para>',  styles["Normal"]))
        
        datahead=[['SPR','Article','Description']]
        
        for c1,c2,c3,c4,c5,c6,c7,c8,c9,c10,c11,c12 in cells:
            if((c1.value=='Closed') and (c5.value=='Fix Submitted') and (str(c9.value)=="Yes") and (str(c10.value)=="True")): #add conditions for include rel notes and is customer spr C1 indicates col1 and so on...
                r  = requests.get("https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c11.value),headers={"User-Agent":"Mozilla/5.0"})
                data = r.text #getting the whole data from the link above.
                soup = BeautifulSoup(data, 'lxml') #parsing the data from the link.
                articlelink="https://www.ptc.com/en/support/article?source=SPR%20Viewer&n="+str(c12.value) #dynamic link for spr no. by just changing the last 6 letters of the link.
                sprlink="https://support.ptc.com/appserver/common/login/ssl/login.jsp?dest=%2Fappserver%2Fcs%2Fview%2Fspr.jsp%3Fn%3D"+str(c11.value)+"&msg=1" #dynamic link for article no. by just changing the some letters of the link.

                
                    
                for h1 in soup.findAll('h1',{"class":'article_solution_title_h1'}):
                    sprcol = c12.value
                    sprcol = '<link href="' + 'https://support.ptc.com/appserver/common/login/ssl/login.jsp?dest=%2Fappserver%2Fcs%2Fview%2Fspr.jsp%3Fn%3D' + str(sprcol) +'&msg=1"><u>' + str(sprcol) + '</u></link>'

                    articlecol = c11.value
                    articlecol = '<link href="' + 'https://www.ptc.com/en/support/article?source=SPR%20Viewer&n=' + str(articlecol) +'"><u>' + str(articlecol) + '</u></link>'
                    cnt=cnt+1
                    
                    data2=([platypus.Paragraph(sprcol, PS('body')), platypus.Paragraph(articlecol, PS('body')), Paragraph((h1.get_text()).strip(),styles['Normal'])])

                    
           
                datahead.append(data2)

                t=Table(datahead,colWidths=(None, None, 120*mm),rowHeights=None, style=None, splitByRow=True,repeatRows=1) #setting col widths for each col first argument represents first col and so on.....
             
        t.setStyle(TableStyle(
        [('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
         ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
         ('FONTSIZE',(0, 0), (-1, 0),11.96),
         ('FONTNAME',(0, 0), (-1, 0),'Arial'),
         ('ALIGN',(0,-1),(-1,-1),'CENTRE'),
         ('VALIGN', (0, 0), (-1, cnt), 'TOP'),
         ('TOPPADDING',(0, 0), (-1, cnt), 5),
         ('BOTTOMPADDING',(0, 0), (-1, cnt), 5),
         ('BACKGROUND', (0, 0), (-1, 0), colors.white)]))

        elements.append(t)
        doc.build(elements,onFirstPage=addPageNumber, onLaterPages=addPageNumber)

        endtime=datetime.now()

      
        print('pdf generated by name:'+releaseno+'.pdf')
        messagebox.showinfo("PDF Generated!", "pdf generated by name:"+releaseno+".pdf")
        print(f'Time taken : {endtime-start_time}')
      

except:
    
    print("Exception occured!")


try:
    import Tkinter as tk
except ImportError:
    import tkinter as tk

try:
    import ttk
    py3 = False
except ImportError:
    import tkinter.ttk as ttk
    py3 = True

import pdfwordgui_support

def vp_start_gui():
    '''Starting point when module is the main routine.'''
    global val, w, root
    root = tk.Tk()
    pdfwordgui_support.set_Tk_var()
    top = Toplevel1 (root)
    pdfwordgui_support.init(root, top)
 
    root.mainloop()

w = None
def create_Toplevel1(root, *args, **kwargs):
    '''Starting point when module is imported by another program.'''
    global w, w_win, rt
    rt = root
    w = tk.Toplevel (root)
    pdfwordgui_support.set_Tk_var()
    top = Toplevel1 (w)
    pdfwordgui_support.init(w, top, *args, **kwargs)
    return (w, top)

def destroy_Toplevel1():
    global w
    w.destroy()
    w = None

class Toplevel1:
    

    
    def __init__(self, top=None):
        '''This class configures and populates the toplevel window.
           top is the toplevel containing window.'''
        _bgcolor = '#d9d9d9'  # X11 color: 'gray85'
        _fgcolor = '#000000'  # X11 color: 'black'
        _compcolor = '#d9d9d9' # X11 color: 'gray85'
        _ana1color = '#d9d9d9' # X11 color: 'gray85'
        _ana2color = '#ececec' # Closest X11 color: 'gray92'

        var1=IntVar()
        var2=IntVar()

        

        top.geometry("375x181+722+121")
        top.title("CPS_Release Code")
        top.configure(background="#d9d9d9")

        self.sellabel = tk.Label(top)
        self.sellabel.place(relx=0.053, rely=0.11, height=21, width=105)
        self.sellabel.configure(background="#d9d9d9")
        self.sellabel.configure(disabledforeground="#a3a3a3")
        self.sellabel.configure(foreground="#000000")
        self.sellabel.configure(text='''CPS File:''')

        self.sheetlabel = tk.Label(top)
        self.sheetlabel.place(relx=0.053, rely=0.331, height=21, width=108)
        self.sheetlabel.configure(background="#d9d9d9")
        self.sheetlabel.configure(disabledforeground="#a3a3a3")
        self.sheetlabel.configure(foreground="#000000")
        self.sheetlabel.configure(text='''Release no:''')

        self.filenameentry = tk.Entry(top)
        self.filenameentry.place(relx=0.347, rely=0.11, height=20, relwidth=0.437)
        self.filenameentry.configure(background="white")
        self.filenameentry.configure(disabledforeground="#a3a3a3")
        self.filenameentry.configure(font="TkFixedFont")
        self.filenameentry.configure(foreground="#000000")
        self.filenameentry.configure(insertbackground="black")

        self.sheetnameentry = tk.Entry(top)
        self.sheetnameentry.place(relx=0.347, rely=0.331, height=20, relwidth=0.437)
        self.sheetnameentry.configure(background="white")
        self.sheetnameentry.configure(disabledforeground="#a3a3a3")
        self.sheetnameentry.configure(font="TkFixedFont")
        self.sheetnameentry.configure(foreground="#000000")
        self.sheetnameentry.configure(insertbackground="black")

       
               
        self.pdfchk = tk.Checkbutton(top,variable=var1)
        self.pdfchk.place(relx=0.347, rely=0.552, relheight=0.138,relwidth=0.131)
        self.pdfchk.configure(activebackground="#ececec")
        self.pdfchk.configure(activeforeground="#000000")
        self.pdfchk.configure(background="#d9d9d9")
        self.pdfchk.configure(disabledforeground="#a3a3a3")
        self.pdfchk.configure(foreground="#000000")
        self.pdfchk.configure(highlightbackground="#d9d9d9")
        self.pdfchk.configure(highlightcolor="black")
        self.pdfchk.configure(justify='left')
        self.pdfchk.configure(text='''PDF''')


        self.wordchk = tk.Checkbutton(top,variable=var2)
        self.wordchk.place(relx=0.613, rely=0.552, relheight=0.138,relwidth=0.152)
        self.wordchk.configure(activebackground="#ececec")
        self.wordchk.configure(activeforeground="#000000")
        self.wordchk.configure(background="#d9d9d9")
        self.wordchk.configure(disabledforeground="#a3a3a3")
        self.wordchk.configure(foreground="#000000")
        self.wordchk.configure(highlightbackground="#d9d9d9")
        self.wordchk.configure(highlightcolor="black")
        self.wordchk.configure(justify='left')
        self.wordchk.configure(text='''Word''')
        


        def checkStatus():
            if(var1.get()):
                print("Please wait patiently, as your pdf file is being generated..")
                exceltopdf(self.filenameentry.get(),self.sheetnameentry.get())
                #exceltopdf("G:/python/New Report Reduced.xlsx","Sheet1")
                #getURLS()
                
            if(var2.get()):
                print("Please wait patiently, as your word file is being generated..")
                #getURLS()
                #fetchCells(self.filenameentry.get(),self.sheetnameentry.get(),self.sheetnameentry.get())
                toWord(self.filenameentry.get(),self.sheetnameentry.get())
                #toWord("G:/python/New Report Reduced.xlsx","Sheet1")
        
        self.generateBtn = tk.Button(top)
        self.generateBtn.place(relx=0.213, rely=0.773, height=24, width=58)
        self.generateBtn.configure(activebackground="#ececec")
        self.generateBtn.configure(activeforeground="#000000")
        self.generateBtn.configure(background="#d9d9d9")
        self.generateBtn.configure(disabledforeground="#a3a3a3")
        self.generateBtn.configure(foreground="#000000")
        self.generateBtn.configure(command=checkStatus)
        self.generateBtn.configure(highlightbackground="#d9d9d9")
        self.generateBtn.configure(highlightcolor="black")
        self.generateBtn.configure(pady="0")
        self.generateBtn.configure(text='''Generate''')

        

        self.cancelBtn = tk.Button(top,command=top.quit)
        self.cancelBtn.place(relx=0.533, rely=0.773, height=24, width=47)
        self.cancelBtn.configure(activebackground="#ececec")
        self.cancelBtn.configure(activeforeground="#000000")
        self.cancelBtn.configure(background="#d9d9d9")
        self.cancelBtn.configure(disabledforeground="#a3a3a3")
        self.cancelBtn.configure(foreground="#000000")
        self.cancelBtn.configure(highlightbackground="#d9d9d9")
        self.cancelBtn.configure(highlightcolor="black")
        self.cancelBtn.configure(pady="0")
        self.cancelBtn.configure(text='''Cancel''')

        
        

        def selectFile():
            self.filenameentry.insert(0,"")
            filename =  tk.filedialog.askopenfilename(initialdir = "/",title = "Select file",filetypes = (("excel files","*.xlsx"),("all files","*.*")))
            self.filenameentry.insert(0,filename)



        self.selfileBtn = tk.Button(top,command=selectFile)
        self.selfileBtn.place(relx=0.827, rely=0.11, height=24, width=60)
        self.selfileBtn.configure(activebackground="#ececec")
        self.selfileBtn.configure(activeforeground="#000000")
        self.selfileBtn.configure(background="#d9d9d9")
        self.selfileBtn.configure(disabledforeground="#a3a3a3")
        self.selfileBtn.configure(foreground="#000000")
        self.selfileBtn.configure(highlightbackground="#d9d9d9")
        self.selfileBtn.configure(highlightcolor="black")
        self.selfileBtn.configure(pady="0")
        self.selfileBtn.configure(text='''Browse''')

        self.Label3 = tk.Label(top)
        self.Label3.place(relx=0.027, rely=0.552, height=21, width=110)
        self.Label3.configure(background="#d9d9d9")
        self.Label3.configure(disabledforeground="#a3a3a3")
        self.Label3.configure(foreground="#000000")
        self.Label3.configure(text='''Select output format:''')
        self.Label3.configure(width=104)





if __name__ == '__main__':
    vp_start_gui()

